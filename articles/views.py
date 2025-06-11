from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.exceptions import PermissionDenied
from rest_framework import viewsets, status
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.filters import SearchFilter
from django_filters.rest_framework import DjangoFilterBackend
from django.utils.translation import gettext as _
from .models import Article, Approval
from .serializers import ArticleSerializer, ApprovalSerializer
from django.contrib.auth import get_user_model
from docx import Document
import os
from .models import FAQ
from .serializers import FAQSerializer
from .serializers import ArticleUploadSerializer


User = get_user_model()


class ArticleViewSet(viewsets.ModelViewSet):
    queryset = Article.objects.all()
    permission_classes = [IsAuthenticated]
    filter_backends = [DjangoFilterBackend, SearchFilter]
    filterset_fields = ['author', 'created_at']
    search_fields = ['title', 'content']
    parser_classes = [MultiPartParser, FormParser]

    def get_serializer_class(self):
        if self.action == 'create':
            return ArticleUploadSerializer
        return ArticleSerializer

    def create(self, request, *args, **kwargs):
        file = request.FILES.get('file')
        if not file:
            return Response({"error": _("No file uploaded")}, status=status.HTTP_400_BAD_REQUEST)

        try:
            doc = Document(file)
            title = doc.paragraphs[0].text.strip() if doc.paragraphs else "Untitled"
            content = "\n".join([p.text for p in doc.paragraphs[1:]])
        except Exception as e:
            return Response({"error": _("Invalid .docx file"), "detail": str(e)}, status=status.HTTP_400_BAD_REQUEST)

        article = Article.objects.create(author=request.user, title=title, content=content)

        monitors = User.objects.filter(is_monitor=True)
        for monitor in monitors:
            Approval.objects.create(article=article, monitor=monitor)

        doc_path = os.path.join('articles', f'{article.id}_{article.title}.docx')
        os.makedirs('articles', exist_ok=True)
        with open(doc_path, 'wb+') as dest:
            for chunk in file.chunks():
                dest.write(chunk)

        return Response({
            "message": _("Article created"),
            "article_id": article.id,
            "doc_path": doc_path
        }, status=status.HTTP_201_CREATED)

    def retrieve(self, request, *args, **kwargs):
        instance = self.get_object()
        instance.view_count += 1
        instance.save(update_fields=['view_count'])
        serializer = self.get_serializer(instance)
        return Response(serializer.data)

    
    def perform_update(self, serializer):
        if self.request.user == self.get_object().author:
            serializer.save()
        else:
            raise PermissionDenied(_("You can only edit your own articles"))

    def perform_destroy(self, instance):
        if self.request.user == instance.author:
            instance.delete()
        else:
            raise PermissionDenied(_("You can only delete your own articles"))

class ApprovalViewSet(viewsets.ModelViewSet):
    queryset = Approval.objects.all()
    serializer_class = ApprovalSerializer
    permission_classes = [IsAuthenticated]

    def perform_update(self, serializer):
        if self.request.user.is_monitor:
            serializer.save()
        else:
            raise PermissionDenied(_("Only monitors can approve articles"))


class FAQ(viewsets.ModelViewSet):
    queryset = FAQ.objects.all()
    serializer_class = FAQSerializer
    permission_classes = [IsAuthenticated]


    